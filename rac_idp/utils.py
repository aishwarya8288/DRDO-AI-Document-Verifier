import os
import json
import re
import mimetypes
import torch
import pytesseract
import docx
import requests
import fitz  # PyMuPDF
import cv2
import matplotlib
import matplotlib.pyplot as plt
import seaborn as sns
import pandas as pd
from PIL import Image
from pdf2image import convert_from_path
from difflib import SequenceMatcher
from django.conf import settings
import transformers
from transformers import AutoProcessor, AutoModelForTokenClassification
from google.cloud import documentai_v1 as documentai
from google.cloud import storage
import subprocess
from google.auth import credentials
from google.oauth2 import service_account
import pdfplumber
import pytesseract
from io import BytesIO
from google.auth.transport.requests import Request
from concurrent.futures import ThreadPoolExecutor

transformers.logging.set_verbosity_error()



# Load LayoutLM model and processor
MODEL_NAME = "microsoft/layoutlmv3-base"
processor = AutoProcessor.from_pretrained(MODEL_NAME, apply_ocr=False)  # Disable internal OCR
model = AutoModelForTokenClassification.from_pretrained(MODEL_NAME)

# Set up Tesseract OCR path
pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'

# Set up Gemini API
GEMINI_API_KEY = "AIzaSyASdSsKECLMv-my69VBSF5ZgyMrCcCy9rg"
GEMINI_API_URL = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-1.5-flash:generateContent?key={GEMINI_API_KEY}"
matplotlib.use('Agg')




def authenticate_with_service_account():
    # Path to your service account JSON key
    key_path = r'C:\Users\aishw\Downloads\IDP\striking-berm-453819-c7-6c6e2252c973.json'
    
    # Authenticate with the service account
    credentials = service_account.Credentials.from_service_account_file(key_path)
    
    # Set up the Google Cloud client with these credentials
    client = documentai.DocumentProcessorServiceClient(credentials=credentials)
    
    return client

def process_document(file_path):
    # Project, location, and processor details
    project_id = "striking-berm-453819-c7"
    location = "us"
    processor_id = "eb9b24fd911523ca"
    
    # Authenticate with service account
    client = authenticate_with_service_account()

    # Construct the full resource name of the processor
    name = f"projects/{project_id}/locations/{location}/processors/{processor_id}"

    # Get the file extension to handle different types
    file_extension = os.path.splitext(file_path)[1].lower()

    # Open the file (image, pdf, docx) and prepare the document
    with open(file_path, "rb") as file:
        file_content = file.read()

    # Determine the mime type based on file extension
    mime_type = ""
    if file_extension == ".pdf":
        mime_type = "application/pdf"
        document = documentai.types.Document(
            content=file_content,
            mime_type=mime_type
        )
    elif file_extension in [".png", ".jpeg", ".jpg"]:
        mime_type = f"image/{file_extension[1:]}"  # "image/png" or "image/jpeg"
        document = documentai.types.Document(
            content=file_content,
            mime_type=mime_type
        )
    elif file_extension == ".docx":
        mime_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        # Convert DOCX to text (simple text extraction approach)
        document_text = extract_text_from_docx(file_path)
        # Return the extracted text as it's not directly processed by Document AI
        return document_text
    else:
        print(f"Unsupported file type: {file_extension}")
        return None

    # Prepare the process request
    request = documentai.types.ProcessRequest(
        name=name,
        raw_document=documentai.types.RawDocument(content=file_content, mime_type=mime_type)
    )

    try:
        # Process the document using the Document AI client
        result = client.process_document(request=request)

        # Extract the document text from the result
        extracted_text = result.document.text

        # Log the extracted text for debugging (optional)
        print("Extracted Text from Document:", extracted_text)

        # Return the extracted text
        return extracted_text

    except Exception as e:
        print(f"Error processing document: {e}")
        return None


def extract_text_from_docx(docx_path):
    """Extract text from DOCX file."""
    try:
        doc = docx.Document(docx_path)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        return '\n'.join(full_text)
    except Exception as e:
        print(f"Error extracting text from DOCX: {e}")
        return None
# Function to process extracted text using Gemini API
  

def send_to_gemini(extracted_text, expected_data):
    """Send extracted text & expected data to Gemini API and return a structured JSON response."""
    try:
        print("üîÑ Sending extracted text & expected data to Gemini API...")

        payload = {
            "contents": [
                {
                    "parts": [
                        {
                            "text": f"""
                            You are an AI that extracts structured data from text.
                            Here is a document's extracted raw text. Your goal is to extract values based on the expected structure. 

                            **Guidelines:**
                            - If a field name is missing in the extracted text, infer it using its value.
                            - If a value is missing, return `null`.
                            - **Ensure valid JSON syntax. Return only JSON, no explanations.**
                            - Do not add extra keys or modify the expected structure.
                            - If a value appears but is not labeled, match it to the closest expected field using semantic similarity.

                            **Extracted Data:**
                            {extracted_text}

                            **Expected Data Format:**
                            {json.dumps(expected_data, indent=2)}

                            **Output JSON:**
                            """
                        }
                    ]
                }
            ]
        }

        headers = {"Content-Type": "application/json"}
        response = requests.post(GEMINI_API_URL, headers=headers, json=payload)

        if response.status_code != 200:
            print(f"‚ùå API Error: {response.status_code} - {response.text}")
            return {"error": f"API Error: {response.status_code}"}

        response_json = response.json()

        # Extract the actual response text from Gemini
        if "candidates" in response_json and len(response_json["candidates"]) > 0:
            gemini_text = response_json["candidates"][0]["content"]["parts"][0]["text"]
            
            try:
                # Direct JSON parsing attempt
                structured_response = json.loads(gemini_text)
                return structured_response
            except json.JSONDecodeError:
                print("‚ö†Ô∏è Raw response is not valid JSON. Attempting extraction...")

                # Extract JSON content from raw response using regex
                json_match = re.search(r'\{.*\}', gemini_text, re.DOTALL)
                if json_match:
                    try:
                        extracted_json = json.loads(json_match.group(0))
                        return extracted_json
                    except json.JSONDecodeError:
                        print("‚ùå Extracted content is still not valid JSON.")

                return {"error": "Invalid JSON response", "raw_response": gemini_text}

        return {"error": "No valid response from Gemini API"}

    except requests.exceptions.RequestException as e:
        print(f"‚ùå Request Exception: {e}")
        return {"error": str(e)}

    



def calculate_mismatch_percentage(expected, extracted):
    """Calculate mismatch percentage between expected and extracted values."""
    if not expected or not extracted:
        return 100.0  # Full mismatch if one of them is missing
    similarity = SequenceMatcher(None, str(expected), str(extracted)).ratio()
    return round((1 - similarity) * 100, 2)  # Convert to percentage

def validate_gemini_response(expected_data, gemini_response):
    """Compares expected data with Gemini's response and returns mismatches & similarities."""
    
    mismatches = {}
    similarities = {}
    missing_fields = []
    extra_fields = []

    # Check for missing fields in Gemini's response
    for key, expected_value in expected_data.items():
        if key not in gemini_response:
            missing_fields.append(key)
        else:
            extracted_value = gemini_response[key]

            # Convert both to lowercase for case-insensitive comparison
            if isinstance(expected_value, str) and isinstance(extracted_value, str):
                expected_value_lower = expected_value.lower().strip()
                extracted_value_lower = extracted_value.lower().strip()
            else:
                expected_value_lower = expected_value
                extracted_value_lower = extracted_value

            mismatch_percentage = calculate_mismatch_percentage(expected_value_lower, extracted_value_lower)

            if mismatch_percentage == 0:
                similarities[key] = extracted_value  # ‚úÖ Exact match
            else:
                mismatches[key] = {
                    "Expected": expected_value,
                    "Extracted": extracted_value,
                    "Mismatch %": mismatch_percentage  # ‚ùå Store mismatch percentage
                }

    # Check for extra fields present in Gemini response but not expected
    for key in gemini_response.keys():
        if key not in expected_data:
            extra_fields.append(key)

    # Structure the final validation report
    validation_result = {
        "mismatches": mismatches,
        "similarities": similarities,
        "missing_fields": missing_fields,
        "extra_fields": extra_fields,
        "mismatch_percentage":mismatch_percentage
    }

    return validation_result




# ‚úÖ Mismatch Visualization Function
def visualize_mismatches(validation_result, applicant_id):
    """
    Generate a mismatch visualization chart based on validation results.
    """
    if not isinstance(validation_result, dict) or not validation_result.get("mismatches"):
        print("‚úÖ No mismatches found, skipping visualization.")
        return None

    mismatches = validation_result["mismatches"]
    mismatch_data = []
    
    try:
        # Extract mismatch data
        for field, values in mismatches.items():
            mismatch_data.append({
                "Field": field,
                "Expected": str(values.get("Expected", "N/A")),  # Correct keys
                "Extracted": str(values.get("Extracted", "N/A")),
                "Mismatch Percentage": values.get("Mismatch %", 0)  # Extract mismatch percentage
            })

        # Ensure we have data to visualize
        if not mismatch_data:
            print("‚úÖ No valid mismatch data to visualize.")
            return None

        # Create DataFrame for visualization
        df = pd.DataFrame(mismatch_data)
        
        if df.empty:
            print("‚úÖ DataFrame is empty, skipping visualization.")
            return None

        # Plot mismatches with correct percentages
        plt.figure(figsize=(12, 6))
        sns.barplot(x="Field", y="Mismatch Percentage", hue="Field", data=df, palette="coolwarm", legend=False)

        plt.xticks(rotation=45, ha='right')
        plt.xlabel("Field Name")
        plt.ylabel("Mismatch Percentage (%)")
        plt.title("Mismatches in Application")

        # Save plot in MEDIA folder
        media_path = settings.MEDIA_ROOT if hasattr(settings, "MEDIA_ROOT") else "media"
        os.makedirs(media_path, exist_ok=True)

        mismatch_plot_path = os.path.join(media_path, f"mismatch_plot_{applicant_id}.png")
        plt.savefig(mismatch_plot_path, bbox_inches="tight")
        plt.close()

        return mismatch_plot_path

    except Exception as e:
        print(f"‚ùå Error in visualize_mismatches: {e}")
        return None